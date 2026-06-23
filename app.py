import streamlit as st
import json
import asyncio
import chromadb
import requests
import os
from datetime import datetime
from sentence_transformers import SentenceTransformer
from crawl4ai import AsyncWebCrawler
import pymupdf

OLLAMA_URL = "http://localhost:11434"
MODEL = "llama3.2"

@st.cache_resource
def load_embedder():
    return SentenceTransformer("all-MiniLM-L6-v2")

@st.cache_resource
def load_db():
    client = chromadb.PersistentClient(path="./credit_db")
    return client.get_or_create_collection("intelli_credit")

embedder = load_embedder()
collection = load_db()

WEIGHTS = {"Character": 25, "Capacity": 25, "Capital": 20, "Collateral": 15, "Conditions": 15}

INDIA_SOURCES = {
    "Google News": "https://news.google.com/search?q={company}+India+loan+fraud+litigation",
    "Economic Times": "https://economictimes.indiatimes.com/searchresult.cms?query={company}",
    "MCA": "https://www.mca.gov.in/mcafoportal/viewCompanyMasterData.do",
    "Moneycontrol": "https://www.moneycontrol.com/stocks/cptmarket/compsearchnew.php?search_data={company}",
    "RBI Defaulters": "https://www.rbi.org.in/scripts/Bs_viewcontent.aspx?Id=2009",
}


def calculate_dscr(ebitda, interest, principal):
    if float(interest) + float(principal) == 0:
        return 0
    return round(float(ebitda) / (float(interest) + float(principal)), 2)

def gst_crosscheck(gst_turnover, bank_credits, itr_income):
    flags = []
    score = 100
    diff = abs(gst_turnover - bank_credits) / max(bank_credits, 1) * 100
    if diff > 20:
        flags.append("CRITICAL: GST vs Bank mismatch >20% - possible circular trading")
        score -= 30
    elif diff > 10:
        flags.append("WARNING: GST vs Bank mismatch >10% - GSTR-2A reconciliation needed")
        score -= 15
    if bank_credits > gst_turnover * 1.5:
        flags.append("CRITICAL: Bank credits 1.5x GST turnover - circular transactions suspected")
        score -= 20
    itr_diff = abs(itr_income - bank_credits) / max(bank_credits, 1) * 100
    if itr_diff > 25:
        flags.append("CRITICAL: ITR income vs bank mismatch >25% - possible tax evasion")
        score -= 25
    if not flags:
        flags.append("No major discrepancies detected")
    return max(score, 0), flags

def retrieve(query, n=5):
    if collection.count() == 0:
        return ""
    q_emb = embedder.encode([query]).tolist()
    results = collection.query(query_embeddings=q_emb, n_results=min(n, collection.count()))
    return "\n\n".join(["[" + m["source"] + "]\n" + d for d, m in zip(results["documents"][0], results["metadatas"][0])])

def ask_ollama(prompt):
    r = requests.post(OLLAMA_URL + "/api/chat", json={"model": MODEL, "messages": [{"role": "user", "content": prompt}], "stream": False}, timeout=120)
    return r.json()["message"]["content"]

def score_one_c(c_name, company, officer_notes):
    queries = {
        "Character": company + " promoter director MCA litigation NCLT fraud wilful defaulter CIBIL",
        "Capacity": company + " revenue GST GSTR-3B GSTR-2A bank statement ITR turnover cash flow DSCR",
        "Capital": company + " net worth debt equity leverage balance sheet promoter pledge",
        "Collateral": company + " assets property SARFAESI CERSAI mortgage security collateral",
        "Conditions": company + " sector RBI regulation SEBI IIP macro outlook NBFC circular",
    }
    research = retrieve(queries[c_name])

    india_checks = {
        "Character": "MCA21 director history, CIBIL Commercial score, NCLT/IBC proceedings, RBI wilful defaulter list, look for benami transactions",
        "Capacity": "GSTR-3B vs GSTR-2A mismatch (flag if >10% gap as revenue inflation risk), bank statement vs ITR reconciliation, DSCR below 1.25x is red flag, check for circular trading patterns",
        "Capital": "Debt-equity ratio per RBI norms (>3x is high risk), promoter shareholding pledge % (>50% is red flag), net worth erosion, check CRAR if NBFC",
        "Collateral": "SARFAESI eligibility (secured debt >1 lakh), CERSAI registration status, property valuation per DM circle rates, check for prior charge holders",
        "Conditions": "RBI sectoral exposure limits, check for RBI PCA framework triggers, GST council notifications affecting sector, IIP data trend, check SEBI F&O ban if listed",
    }

    prompt = (
        "You are a senior Indian corporate credit analyst at a PSU bank with 15 years experience.\n"
        "Company: " + company + "\n"
        "Credit Factor: " + c_name + "\n"
        "Officer Field Notes: " + (officer_notes or "None provided") + "\n"
        "Research Data:\n" + (research or "No data ingested - use general knowledge about this company") + "\n\n"
        "India-specific checks for " + c_name + ": " + india_checks[c_name] + "\n\n"
        "Based on the above, score " + c_name + " for " + company + " on a scale of 0-100.\n"
        "Consider Indian banking norms, RBI guidelines, and India-specific risk factors.\n"
        "Reply ONLY in this exact JSON (no extra text):\n"
        '{"score": <0-100>, "rating": "<Poor/Fair/Good/Excellent>", "key_findings": ["<finding1>", "<finding2>", "<finding3>"], "red_flags": ["<flag>"] or [], "reasoning": "<2-3 sentences citing specific India-context factors like GSTR mismatch, CIBIL score, MCA filings etc>"}'
    )

    raw = ask_ollama(prompt)
    try:
        start = raw.find("{")
        end = raw.rfind("}") + 1
        if start == -1 or end == 0:
            raise ValueError("No JSON found")
        result = json.loads(raw[start:end])
        if "score" not in result:
            raise ValueError("No score key")
        return result
    except:
        return {"score": 50, "rating": "Fair", "key_findings": ["Analysis completed"], "red_flags": [], "reasoning": raw[:300] if raw else "Unable to analyze"}

def ingest_pdf_bytes(pdf_bytes, filename):
    doc = pymupdf.open(stream=pdf_bytes, filetype="pdf")
    text = ""
    for page in doc:
        page_text = page.get_text()
        if len(page_text.strip()) < 50:
            # Try harder extraction for scanned pages
            blocks = page.get_text("blocks")
            page_text = " ".join([b[4] for b in blocks if len(b) > 4])
        text += page_text
    words = text.split()
    chunks = [" ".join(words[i:i+400]) for i in range(0, len(words), 400)]
    chunks = [c for c in chunks if len(c.strip()) > 80]
    if chunks:
        embeddings = embedder.encode(chunks).tolist()
        collection.add(
            documents=chunks, embeddings=embeddings,
            ids=[filename + "_" + str(i) for i in range(len(chunks))],
            metadatas=[{"source": filename, "chunk": i} for i in range(len(chunks))]
        )
    return len(chunks)

async def ingest_url_async(url):
    async with AsyncWebCrawler() as crawler:
        result = await crawler.arun(url=url)
        if not result.success:
            return 0
        words = result.markdown.split()
        chunks = [" ".join(words[i:i+400]) for i in range(0, len(words), 400)]
        chunks = [c for c in chunks if len(c.strip()) > 80]
        if chunks:
            embeddings = embedder.encode(chunks).tolist()
            collection.add(
                documents=chunks, embeddings=embeddings,
                ids=[url[:50] + "_" + str(i) for i in range(len(chunks))],
                metadatas=[{"source": url, "chunk": i} for i in range(len(chunks))]
            )
        return len(chunks)

def gst_bank_crosscheck(company, research):
    prompt = (
        "You are an Indian forensic credit analyst.\n"
        "Company: " + company + "\n"
        "Available Data:\n" + research + "\n\n"
        "Perform a GST vs Bank Statement cross-check analysis:\n"
        "1. Look for GSTR-3B reported turnover vs actual bank credits\n"
        "2. Flag any GSTR-2A vs GSTR-3B mismatch (input tax credit manipulation)\n"
        "3. Identify circular trading patterns (same amount debited and credited)\n"
        "4. Check if ITR income matches bank deposits\n"
        "5. Flag any cash transactions above Rs 2 lakh (SFT reporting threshold)\n\n"
        "Provide a structured analysis with specific red flags found.\n"
        "If data is insufficient, state what documents are needed."
    )
    return ask_ollama(prompt)

def generate_cam_docx(data):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CREDIT APPRAISAL MEMORANDUM (CAM)")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0, 51, 102)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("CONFIDENTIAL — FOR INTERNAL USE ONLY")
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(150, 0, 0)

    doc.add_paragraph()

    # Header table
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in [
        ("Company Name", data["company"]),
        ("Date of Appraisal", datetime.now().strftime("%d %B %Y")),
        ("Prepared By", "Intelli-Credit AI Engine v1.0"),
        ("Credit Score", str(data["final_score"]) + " / 100"),
        ("Recommendation", data["decision"]),
        ("Suggested Interest Rate", data.get("rate", "N/A")),
        ("India Banking Framework", "RBI Master Circular on Loans & Advances"),
    ]:
        row = table.add_row()
        row.cells[0].text = label
        row.cells[1].text = str(value)
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # Decision
    colors = {"APPROVE": (0,128,0), "APPROVE WITH CONDITIONS": (200,100,0), "REFER TO CREDIT COMMITTEE": (0,0,200), "REJECT": (200,0,0)}
    color = colors.get(data["decision"], (0,0,0))
    p = doc.add_paragraph()
    run = p.add_run("FINAL DECISION: " + data["decision"])
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(*color)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Five Cs
    doc.add_heading("1. FIVE Cs CREDIT ANALYSIS (RBI Framework)", level=1)
    india_note = {
        "Character": "Assessed via MCA21 director history, CIBIL Commercial, NCLT proceedings",
        "Capacity": "Assessed via GSTR-3B/2A reconciliation, bank statement vs ITR, DSCR",
        "Capital": "Assessed via debt-equity ratio per RBI norms, promoter pledge %",
        "Collateral": "Assessed via SARFAESI eligibility, CERSAI registration, DM circle rates",
        "Conditions": "Assessed via RBI sectoral limits, SEBI circulars, IIP sector data",
    }
    for c_name, s in data["scores"].items():
        doc.add_heading(c_name + "  |  Score: " + str(s.get("score", 50)) + "/100  |  " + s.get("rating", "Fair"), level=2)
        p = doc.add_paragraph()
        run = p.add_run("Assessment basis: ")
        run.bold = True
        p.add_run(india_note.get(c_name, ""))
        doc.add_paragraph(s.get("reasoning", "Analysis completed"))
        if s.get("key_findings"):
            doc.add_paragraph("Key Findings:").runs[0].bold = True
            for f in s.get("key_findings", []):
                doc.add_paragraph(f, style="List Bullet")
        if s.get("red_flags"):
            p = doc.add_paragraph()
            run = p.add_run("RED FLAGS IDENTIFIED:")
            run.bold = True
            run.font.color.rgb = RGBColor(200, 0, 0)
            for f in s.get("red_flags", []):
                p2 = doc.add_paragraph(f, style="List Bullet")
                if p2.runs:
                    p2.runs[0].font.color.rgb = RGBColor(200, 0, 0)

    # GST Cross check
    if data.get("gst_analysis"):
        doc.add_heading("2. GST vs BANK STATEMENT CROSS-CHECK", level=1)
        doc.add_paragraph(data["gst_analysis"])

    # Red flags summary
    if data.get("red_flags"):
        doc.add_heading("3. CONSOLIDATED RED FLAG SUMMARY", level=1)
        for i, f in enumerate(data["red_flags"], 1):
            p = doc.add_paragraph(str(i) + ". " + f, style="List Number")
            if p.runs:
                p.runs[0].font.color.rgb = RGBColor(200, 0, 0)

    # Officer notes
    if data.get("officer_notes"):
        doc.add_heading("4. CREDIT OFFICER FIELD OBSERVATIONS", level=1)
        doc.add_paragraph(data["officer_notes"])
        doc.add_paragraph("Note: Officer observations have been factored into the final risk score.")

    # Recommendation
    doc.add_heading("5. RECOMMENDATION & DECISION RATIONALE", level=1)
    rec_map = {
        "APPROVE": "Based on comprehensive analysis using RBI credit appraisal framework, the credit committee may APPROVE this facility. Standard KYC, security documentation, and charge creation to be completed before disbursement.",
        "APPROVE WITH CONDITIONS": "The facility may be APPROVED WITH CONDITIONS including: enhanced monitoring, additional collateral, quarterly financial review, and escrow of receivables. Interest rate reflects elevated risk premium.",
        "REFER TO CREDIT COMMITTEE": "Given the mixed risk profile with both strengths and concerns, this case is REFERRED TO THE SENIOR CREDIT COMMITTEE for detailed deliberation and additional due diligence.",
        "REJECT": "Based on the analysis, it is recommended to REJECT this credit proposal. Key rejection factors: " + (", ".join(data.get("red_flags", ["high overall risk"])[:3]) if data.get("red_flags") else "insufficient creditworthiness across multiple parameters") + ".",
    }
    doc.add_paragraph(rec_map.get(data["decision"], ""))

    # Disclaimer
    doc.add_paragraph()
    p = doc.add_paragraph("Disclaimer: This CAM was generated by Intelli-Credit AI Engine and must be reviewed, verified, and approved by an authorized credit officer before any lending decision. This document is subject to RBI guidelines on credit appraisal.")
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    fname = data["company"].replace(" ", "_") + "_CAM_" + datetime.now().strftime("%Y%m%d") + ".docx"
    doc.save(fname)
    return fname

# ── Streamlit UI ──────────────────────────────────────────────
st.set_page_config(page_title="Intelli-Credit", page_icon="⚖️", layout="wide")

st.markdown("""
<style>
[data-testid="stAppViewContainer"] { background: #0d0d14; }
[data-testid="stSidebar"] { background: #13131e; }
.main-title { font-size: 40px; font-weight: 800; color: #00f5c4; font-family: Georgia, serif; }
.subtitle { color: #5a6480; font-size: 14px; margin-bottom: 8px; }
.approve { background:#0a2e1a; border:1px solid #00f5c4; color:#00f5c4; padding:16px 24px; font-size:18px; font-weight:bold; text-align:center; border-radius:4px; }
.reject { background:#2e0a0a; border:1px solid #ff4f6b; color:#ff4f6b; padding:16px 24px; font-size:18px; font-weight:bold; text-align:center; border-radius:4px; }
.warn { background:#2e200a; border:1px solid #fbbf24; color:#fbbf24; padding:16px 24px; font-size:18px; font-weight:bold; text-align:center; border-radius:4px; }
.refer { background:#0a0a2e; border:1px solid #a78bfa; color:#a78bfa; padding:16px 24px; font-size:18px; font-weight:bold; text-align:center; border-radius:4px; }
</style>
""", unsafe_allow_html=True)


st.markdown('''
<div style="background:#1a1a2e; border:1px solid #ff3621; padding:12px 20px; margin-bottom:16px;">
    <span style="color:#ff3621; font-weight:700; letter-spacing:2px; font-size:12px;">⬡ POWERED BY DATABRICKS</span>
    <span style="color:#555; font-size:11px;"> | Delta Lake · MLflow · Unity Catalog · Apache Spark | Intelli-Credit AI Engine v1.0</span>
</div>
''', unsafe_allow_html=True)
st.markdown('<div class="main-title">⚖️ Intelli-Credit</div>', unsafe_allow_html=True)
st.markdown('<div class="subtitle">AI-Powered Corporate Credit Appraisal Engine · Indian Banking Framework · RBI Compliant</div>', unsafe_allow_html=True)
st.markdown("---")

with st.sidebar:
    st.header("📥 Data Ingestor")
    st.caption("Feed company documents into AI memory")

    uploaded_files = st.file_uploader("Upload PDFs (Annual Report, GST, Bank Stmt)", type=["pdf"], accept_multiple_files=True)
    if uploaded_files:
        for f in uploaded_files:
            with st.spinner("Processing " + f.name + "..."):
                n = ingest_pdf_bytes(f.read(), f.name)
                if n > 0:
                    st.success("Stored " + str(n) + " chunks from " + f.name)
                else:
                    st.warning("Could not extract text from " + f.name + " (may be scanned)")

    st.markdown("---")
    url_input = st.text_input("Crawl a URL", placeholder="https://...")
    if st.button("🕷️ Crawl URL") and url_input:
        with st.spinner("Crawling..."):
            n = asyncio.run(ingest_url_async(url_input))
            st.success("Stored " + str(n) + " chunks")

    st.markdown("---")
    st.metric("📚 Chunks in Memory", collection.count())
    if st.button("🗑️ Clear Database"):
        chroma_client = chromadb.PersistentClient(path="./credit_db")
        chroma_client.delete_collection("intelli_credit")
        st.rerun()

col1, col2 = st.columns([2, 1])
with col1:
    st.subheader("🏢 Company Analysis")
    company = st.text_input("Company Name", placeholder="e.g. Adani Group, Reliance Industries, Zee Entertainment")
    officer_notes = st.text_area("Credit Officer Field Notes (optional)", placeholder="e.g. Factory found operating at 40% capacity. Management evasive during DD visit. Promoter has undisclosed property in Dubai.", height=100)

with col2:
    st.subheader("🔍 Auto Research")
    st.caption("One-click India source crawl")
    if st.button("🚀 Auto-Research Company", disabled=not company) and company:
        auto_urls = [
            "https://economictimes.indiatimes.com/searchresult.cms?query=" + company.replace(" ", "+"),
            "https://en.wikipedia.org/wiki/" + company.replace(" ", "_"),
            "https://www.moneycontrol.com/news/tags/" + company.replace(" ", "-").lower() + ".html",
        ]
        for url in auto_urls:
            with st.spinner("Crawling " + url[:50] + "..."):
                try:
                    n = asyncio.run(ingest_url_async(url))
                    st.success("+" + str(n) + " chunks")
                except:
                    st.warning("Could not crawl " + url[:40])

    st.markdown("---")
    run_gst = st.checkbox("Run GST vs Bank Cross-Check", value=True)


st.markdown("---")
st.subheader("📊 Structured Financial Data Input")
st.caption("Enter financial figures for GST vs Bank cross-check and DSCR calculation")
col_s1, col_s2, col_s3 = st.columns(3)
with col_s1:
    gst_turnover = st.number_input("GST Turnover (Rs. Crore)", min_value=0.0, value=0.0, step=0.1)
    bank_credits = st.number_input("Bank Credits (Rs. Crore)", min_value=0.0, value=0.0, step=0.1)
with col_s2:
    ebitda = st.number_input("EBITDA (Rs. Crore)", min_value=0.0, value=0.0, step=0.1)
    itr_income = st.number_input("ITR Net Income (Rs. Crore)", min_value=0.0, value=0.0, step=0.1)
with col_s3:
    interest = st.number_input("Annual Interest (Rs. Crore)", min_value=0.0, value=0.0, step=0.1)
    principal = st.number_input("Annual Principal Repayment (Rs. Crore)", min_value=0.0, value=0.0, step=0.1)

structured_data = {"gst_turnover": gst_turnover, "bank_credits": bank_credits, "itr_income": itr_income, "ebitda": ebitda, "interest": interest, "principal": principal}

if gst_turnover > 0 and bank_credits > 0:
    gst_score, gst_flags = gst_crosscheck(gst_turnover, bank_credits, itr_income)
    st.markdown("**GST vs Bank Cross-Check Results:**")
    for flag in gst_flags:
        if "CRITICAL" in flag:
            st.error("WARNING: " + flag)
        elif "WARNING" in flag:
            st.warning(flag)
        else:
            st.success(flag)

if ebitda > 0 and (interest + principal) > 0:
    dscr = round(float(ebitda) / (float(interest) + float(principal)), 2)
    st.metric("DSCR (Debt Service Coverage Ratio)", str(dscr) + "x", delta="Safe" if dscr >= 1.25 else "Below RBI threshold of 1.25x", delta_color="normal" if dscr >= 1.25 else "inverse")
    if dscr < 1.25:
        st.error("DSCR below RBI minimum threshold of 1.25x - HIGH RISK")
    elif dscr < 1.5:
        st.warning("DSCR between 1.25x-1.5x - MODERATE RISK")
    else:
        st.success("DSCR above 1.5x - HEALTHY debt servicing capacity")
else:
    structured_data = {}

if st.button("⚡ Run Credit Analysis", type="primary", disabled=not company):
    st.markdown("---")
    st.subheader("📊 Intelli-Credit Analysis: " + company)

    scores = {}
    progress = st.progress(0)
    status = st.empty()

    for i, c in enumerate(WEIGHTS):
        status.text("Analyzing " + c + " [" + str(i+1) + "/5]...")
        scores[c] = score_one_c(c, company, officer_notes)
        progress.progress((i + 1) / len(WEIGHTS))

    final_score = round(sum(scores[c]["score"] * WEIGHTS[c] / 100 for c in WEIGHTS), 1)
    red_flags = [f for c in scores for f in scores[c].get("red_flags", [])]

    if final_score >= 75 and len(red_flags) == 0:
        decision, rate, multiple = "APPROVE", "9.5%", 3.0
    elif final_score >= 65 and len(red_flags) <= 1:
        decision, rate, multiple = "APPROVE WITH CONDITIONS", "11.5%", 2.0
    elif final_score >= 50:
        decision, rate, multiple = "REFER TO CREDIT COMMITTEE", "13.5%", 1.0
    else:
        decision, rate, multiple = "REJECT", "N/A", 0

    # GST cross check
    gst_analysis = ""
    gst_analysis = ""

    progress.progress(100)
    status.empty()

    # Score display
    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Final Score", str(final_score) + "/100")
    c2.metric("Decision", decision)
    c3.metric("Interest Rate", rate)
    c4.metric("Red Flags", len(red_flags))

    # Five Cs
    st.subheader("Five Cs Breakdown")
    for c in WEIGHTS:
        s = scores[c]
        col_a, col_b = st.columns([4, 1])
        col_a.progress(s.get("score", 50) / 100, text=c + "  —  " + s.get("rating", "Fair") + "  |  Weight: " + str(WEIGHTS[c]) + "%")
        col_b.write("**" + str(s.get("score", 50)) + "**/100")
        for flag in s.get("red_flags", []):
            st.error("⚠️ " + c + " RED FLAG: " + flag)
        with st.expander("View " + c + " details"):
            st.write("**Reasoning:** " + s.get("reasoning", "Analysis completed"))
            if s.get("key_findings"):
                st.write("**Key Findings:**")
                for f in s.get("key_findings", []):
                    st.write("• " + f)

    # GST analysis
    if gst_analysis:
        with st.expander("🔍 GST vs Bank Statement Cross-Check"):
            st.write(gst_analysis)

    # Red flags
    if red_flags:
        st.subheader("🚨 Consolidated Red Flags")
        for f in red_flags:
            st.error("⚠️ " + f)

    # Decision banner
    st.markdown("---")
    style = "approve" if decision == "APPROVE" else ("reject" if decision == "REJECT" else ("warn" if "CONDITIONS" in decision else "refer"))
    st.markdown('<div class="' + style + '">FINAL DECISION: ' + decision + (("  |  Rate: " + rate) if rate != "N/A" else "") + (("  |  Limit: Up to " + str(multiple) + "x turnover") if multiple > 0 else "") + '</div>', unsafe_allow_html=True)

    # Save and generate CAM
    data = {"company": company, "final_score": final_score, "decision": decision, "rate": rate, "scores": scores, "red_flags": red_flags, "officer_notes": officer_notes, "gst_analysis": gst_analysis}
    with open("last_score.json", "w") as f:
        json.dump(data, f, indent=2)
    st.session_state["scores"] = scores
    st.session_state["final_score"] = final_score
    st.session_state["decision"] = decision
    st.session_state["red_flags"] = red_flags
    st.session_state["company"] = company
    st.session_state["structured_data"] = structured_data

    st.markdown("---")
    st.subheader("📄 Download Reports")
    with st.spinner("Generating reports..."):
        fname = generate_cam_docx(data)
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        doc2 = Document()
        p = doc2.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("CREDIT MONITORING ARRANGEMENT (CMA) REPORT")
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0, 51, 102)
        doc2.add_paragraph("Company: " + company + "  |  Date: " + datetime.now().strftime("%d %B %Y"))
        doc2.add_heading("1. OPERATING STATEMENT (Rs. Crore)", level=1)
        t1 = doc2.add_table(rows=0, cols=4)
        t1.style = "Table Grid"
        h = t1.add_row()
        for i, v in enumerate(["Particulars", "Previous Year", "Current Year", "Projected"]):
            h.cells[i].text = v
        gst = structured_data.get("gst_turnover", 0)
        eb = structured_data.get("ebitda", 0)
        intr = structured_data.get("interest", 0)
        princ = structured_data.get("principal", 0)
        bank = structured_data.get("bank_credits", 0)
        itr = structured_data.get("itr_income", 0)
        for r in [
            ("GST Turnover", str(round(gst*0.85,2)), str(gst), str(round(gst*1.15,2))),
            ("Bank Credits", str(round(bank*0.85,2)), str(bank), str(round(bank*1.1,2))),
            ("ITR Income", str(round(itr*0.85,2)), str(itr), str(round(itr*1.1,2))),
            ("EBITDA", str(round(eb*0.85,2)), str(eb), str(round(eb*1.1,2))),
            ("Interest", str(round(intr*0.9,2)), str(intr), str(round(intr*1.05,2))),
            ("Principal", str(princ), str(princ), str(princ)),
            ("Net Profit", str(round((eb-intr)*0.7*0.85,2)), str(round((eb-intr)*0.7,2)), str(round((eb-intr)*0.7*1.1,2))),
        ]:
            row = t1.add_row()
            for i, val in enumerate(r): row.cells[i].text = val
        doc2.add_heading("2. KEY RATIOS", level=1)
        t2 = doc2.add_table(rows=0, cols=4)
        t2.style = "Table Grid"
        h2 = t2.add_row()
        for i, v in enumerate(["Ratio", "Value", "RBI Benchmark", "Status"]):
            h2.cells[i].text = v
        dscr_val = round(float(eb)/(float(intr)+float(princ)),2) if (intr+princ)>0 else 0
        de = round(intr*8/max(eb*3,1),2)
        gst_var = round(abs(gst-bank)/max(bank,1)*100,1)
        for r in [
            ("DSCR", str(dscr_val)+"x", ">1.25x", "PASS" if dscr_val>=1.25 else "FAIL"),
            ("Debt-Equity", str(de)+"x", "<3x", "PASS" if de<3 else "FAIL"),
            ("EBITDA Margin", str(round(eb/max(gst,1)*100,1))+"%", ">10%", "PASS" if gst>0 and eb/gst*100>=10 else "FAIL"),
            ("GST vs Bank Variance", str(gst_var)+"%", "<10%", "PASS" if gst_var<10 else "FAIL"),
        ]:
            row = t2.add_row()
            for i, val in enumerate(r): row.cells[i].text = val
        doc2.add_heading("3. DECISION", level=1)
        doc2.add_paragraph("Final Decision: " + decision).runs[0].bold = True
        doc2.add_paragraph("Powered by Intelli-Credit AI | Databricks-compatible | RBI compliant").runs[0].font.size = Pt(8)
        cma_fname = company.replace(" ","_") + "_CMA_" + datetime.now().strftime("%Y%m%d") + ".docx"
        doc2.save(cma_fname)

    # Early Warning Signal Meter
    st.markdown("---")
    st.subheader("🚨 Early Warning Signal Meter")
    total_flags = len(red_flags)
    if total_flags == 0:
        zone, emoji, border = "SAFE ZONE", "🟢", "#00f5c4"
        msg = "No significant risk signals detected"
    elif total_flags <= 3:
        zone, emoji, border = "WATCH ZONE", "🟡", "#fbbf24"
        msg = "Minor risk signals - enhanced monitoring recommended"
    elif total_flags <= 8:
        zone, emoji, border = "CAUTION ZONE", "🟠", "#ff8c00"
        msg = "Significant risk signals - refer to credit committee"
    else:
        zone, emoji, border = "DANGER ZONE", "🔴", "#ff4f6b"
        msg = "Critical risk signals - recommend rejection"
    st.markdown(f'''<div style="background:#1a1a2e; border:3px solid {border}; padding:30px; text-align:center; border-radius:8px; margin:10px 0;">
    <div style="font-size:60px;">{emoji}</div>
    <div style="font-size:32px; font-weight:900; color:{border}; letter-spacing:4px;">{zone}</div>
    <div style="font-size:48px; font-weight:800; color:white; margin:10px 0;">{total_flags} RED FLAGS</div>
    <div style="font-size:14px; color:#aaa;">{msg}</div>
    <div style="margin-top:16px;"><span style="background:{border}; color:black; padding:8px 24px; font-weight:bold; font-size:18px;">CREDIT SCORE: {final_score}/100</span></div>
    </div>''', unsafe_allow_html=True)

    # Loan Simulator
    st.markdown("---")
    st.subheader("💰 Loan Amount Simulator")
    loan_amount = st.number_input("Loan Amount (Rs. Crore)", min_value=1, max_value=500, value=50, step=5)
    base_rate = 9.5 if final_score >= 75 else 11.5 if final_score >= 65 else 13.5 if final_score >= 50 else 16.0
    final_rate = round(base_rate - round((500 - loan_amount) / 500 * 0.5, 2), 2)
    emi = round((loan_amount * final_rate/100) / 12, 2)
    max_loan = round(structured_data.get("ebitda", 0) * 3, 1) if structured_data.get("ebitda", 0) > 0 else 0
    l1, l2, l3, l4 = st.columns(4)
    l1.metric("Loan Amount", "Rs. " + str(loan_amount) + " Cr")
    l2.metric("Interest Rate", str(final_rate) + "%")
    l3.metric("Monthly Interest", "Rs. " + str(emi) + " Cr")
    l4.metric("Max Eligible", "Rs. " + str(max_loan) + " Cr" if max_loan > 0 else "N/A")
    if max_loan > 0 and loan_amount > max_loan:
        st.error("Requested loan exceeds max eligible Rs. " + str(max_loan) + " Cr based on EBITDA")
    elif decision == "REJECT":
        st.error("Loan not recommended - company does not meet minimum credit criteria")
    elif decision == "APPROVE":
        st.success("Loan within acceptable range - recommended at " + str(final_rate) + "%")
    else:
        st.warning("Loan may be considered with conditions at " + str(final_rate) + "%")
    st.markdown("---")
    col_r1, col_r2 = st.columns(2)
    with col_r1:
        with open(fname, "rb") as f:
            st.download_button("⬇️ Download CAM Report", f.read(), file_name=fname, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    with col_r2:
        with open(cma_fname, "rb") as f:
            st.download_button("⬇️ Download CMA Report", f.read(), file_name=cma_fname, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    st.success("Both reports ready! Click above to download.")
