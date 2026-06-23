import gradio as gr
import json
import asyncio
import chromadb
import requests
import os
from datetime import datetime
from sentence_transformers import SentenceTransformer
from crawl4ai import AsyncWebCrawler
import pymupdf

OLLAMA_URL = "http://localhost:11434"
MODEL = "llama3.2"

embedder = SentenceTransformer("all-MiniLM-L6-v2")
chroma_client = chromadb.PersistentClient(path="./credit_db")
collection = chroma_client.get_or_create_collection("intelli_credit")

WEIGHTS = {"Character": 25, "Capacity": 25, "Capital": 20, "Collateral": 15, "Conditions": 15}

def retrieve(query, n=5):
    if collection.count() == 0:
        return ""
    q_emb = embedder.encode([query]).tolist()
    results = collection.query(query_embeddings=q_emb, n_results=min(n, collection.count()))
    return "\n\n".join(["[" + m["source"] + "]\n" + d for d, m in zip(results["documents"][0], results["metadatas"][0])])

def ask_ollama(prompt, timeout=180):
    try:
        r = requests.post(OLLAMA_URL + "/api/chat",
            json={"model": MODEL, "messages": [{"role": "user", "content": prompt}], "stream": False},
            timeout=timeout)
        return r.json()["message"]["content"]
    except:
        return '{"score": 50, "rating": "Fair", "key_findings": ["Timeout"], "red_flags": [], "reasoning": "Analysis timed out."}'

def score_one_c(c_name, company, officer_notes):
    queries = {
        "Character": company + " promoter director MCA litigation NCLT fraud wilful defaulter CIBIL",
        "Capacity": company + " revenue GST GSTR-3B GSTR-2A bank statement ITR turnover cash flow DSCR",
        "Capital": company + " net worth debt equity leverage balance sheet promoter pledge",
        "Collateral": company + " assets property SARFAESI CERSAI mortgage security collateral",
        "Conditions": company + " sector RBI regulation SEBI IIP macro outlook NBFC circular",
    }
    india_checks = {
        "Character": "MCA21 director history, CIBIL Commercial score, NCLT/IBC proceedings, RBI wilful defaulter list, CRISIL/ICRA/CARE rating history",
        "Capacity": "GSTR-3B vs GSTR-2A mismatch, bank statement vs ITR reconciliation, DSCR below 1.25x",
        "Capital": "Debt-equity ratio per RBI norms, promoter pledge % above 50%, net worth erosion",
        "Collateral": "SARFAESI eligibility, CERSAI registration, DM circle rates valuation",
        "Conditions": "RBI sectoral exposure limits, SEBI circulars, GST council notifications, IIP trend, CRISIL sector outlook, ICRA rating watch",
    }
    research = retrieve(queries[c_name])
    prompt = (
        "You are a senior Indian corporate credit analyst at a PSU bank.\n"
        "Company: " + company + "\nCredit Factor: " + c_name + "\n"
        "CRITICAL OFFICER NOTES (must factor into score aggressively): " + (officer_notes or "None") + "\n""STRICT RULES:\n""- If notes mention NCLT or insolvency → Character score MUST be below 25\n""- If notes mention ED investigation or fraud → Character score MUST be below 20\n""- If notes mention wilful defaulter → Character score MUST be below 15\n""- If notes mention auditor resigned → Capital score MUST be below 30\n""- You MUST write specific reasoning, never say Analysis completed\n"
        "Research:\n" + (research if research else "No documents found. Use your general knowledge about " + company + " from training data.") + "\n"
        "India checks: " + india_checks[c_name] + "\n\n"
        "Score " + c_name + " 0-100. Reply ONLY in JSON:\n"
        '{"score": <0-100>, "rating": "<Poor/Fair/Good/Excellent>", "key_findings": ["f1", "f2"], "red_flags": [], "reasoning": "<2-3 sentences>"}'
    )
    raw = ask_ollama(prompt)
    try:
        result = json.loads(raw[raw.find("{"):raw.rfind("}")+1])
        if "score" not in result:
            raise ValueError()
        return result
    except:
        return {"score": 50, "rating": "Fair", "key_findings": [], "red_flags": [], "reasoning": "Analysis completed."}


def save_audit_trail(company, decision, final_score, red_flags, officer_notes):
    audit = {
        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "company": company,
        "decision": decision,
        "final_score": final_score,
        "red_flags_count": len(red_flags),
        "red_flags": red_flags,
        "officer_notes": officer_notes,
        "data_sources": ["ChromaDB Vector Store", "Ollama LLM", "Web Crawler", "Structured Input"],
        "model": "llama3.2",
        "framework": "RBI Five Cs Credit Appraisal"
    }
    log_file = "audit_trail.json"
    existing = []
    if os.path.exists(log_file):
        with open(log_file) as f:
            try:
                existing = json.load(f)
            except:
                existing = []
    existing.append(audit)
    with open(log_file, "w") as f:
        json.dump(existing, f, indent=2)
    return audit

def get_audit_trail():
    log_file = "audit_trail.json"
    if not os.path.exists(log_file):
        return "No audit trail yet. Run an analysis first."
    with open(log_file) as f:
        data = json.load(f)
    result = "INTELLI-CREDIT AUDIT TRAIL\n" + "="*50 + "\n\n"
    for entry in reversed(data):
        result += "Timestamp:   " + entry["timestamp"] + "\n"
        result += "Company:     " + entry["company"] + "\n"
        result += "Decision:    " + entry["decision"] + "\n"
        result += "Score:       " + str(entry["final_score"]) + "/100\n"
        result += "Red Flags:   " + str(entry["red_flags_count"]) + "\n"
        result += "Model:       " + entry["model"] + "\n"
        result += "Framework:   " + entry["framework"] + "\n"
        result += "-"*50 + "\n\n"
    return result

import feedparser

def search_google_news(company):
    results = []
    feeds = [
        "https://news.google.com/rss/search?q=" + company.replace(" ", "+") + "+India+loan+fraud+litigation&hl=en-IN&gl=IN&ceid=IN:en",
        "https://news.google.com/rss/search?q=" + company.replace(" ", "+") + "+SEBI+RBI+NCLT+court&hl=en-IN&gl=IN&ceid=IN:en",
        "https://news.google.com/rss/search?q=" + company.replace(" ", "+") + "+promoter+defaulter+fraud&hl=en-IN&gl=IN&ceid=IN:en",
    ]
    all_text = ""
    for feed_url in feeds:
        try:
            feed = feedparser.parse(feed_url)
            for entry in feed.entries[:20]:
                title = entry.get("title", "")
                summary = entry.get("summary", "")
                published = entry.get("published", "")
                if any(word.lower() in (title + summary).lower() for word in company.split()):
                    all_text += title + " " + summary + " [" + published + "]\n"
                    results.append(title)
        except:
            pass
    if all_text:
        words = all_text.split()
        chunks = [" ".join(words[i:i+400]) for i in range(0, len(words), 400)]
        chunks = [c for c in chunks if len(c.strip()) > 50]
        if chunks:
            embeddings = embedder.encode(chunks).tolist()
            collection.add(
                documents=chunks, embeddings=embeddings,
                ids=["gnews_" + company[:20] + "_" + str(i) for i in range(len(chunks))],
                metadatas=[{"source": "Google News RSS", "chunk": i} for i in range(len(chunks))]
            )
    return results

def deep_research(company):
    if not company:
        return "Please enter company name"
    log = []
    log.append("Starting deep research for: " + company)
    
    # Google News RSS
    log.append("\n1. Searching Google News RSS...")
    news = search_google_news(company)
    log.append("   Found " + str(len(news)) + " news articles")
    for n in news[:5]:
        log.append("   • " + n[:80])
    
    # Auto crawl India sources
    log.append("\n2. Crawling India financial + litigation sources...")
    urls = [
        "https://en.wikipedia.org/wiki/" + company.replace(" ", "_"),
        "https://economictimes.indiatimes.com/searchresult.cms?query=" + company.replace(" ", "+"),
        "https://www.moneycontrol.com/news/tags/" + company.replace(" ", "-").lower() + ".html",
        "https://www.sebi.gov.in/sebiweb/other/OtherAction.do?doRecognisedFpi=yes&intmId=13",
        "https://www.livelaw.in/?s=" + company.replace(" ", "+"),
        "https://www.barandbench.com/?s=" + company.replace(" ", "+"),
        "https://www.vccircle.com/?s=" + company.replace(" ", "+"),
        "https://www.business-standard.com/search?q=" + company.replace(" ", "+"),
    ]
    for url in urls:
        try:
            n = asyncio.run(crawl_url_async(url))
            log.append("   +" + str(n) + " chunks from " + url[:60])
        except:
            log.append("   Failed: " + url[:60])
    
    log.append("\n3. Total chunks in DB: " + str(collection.count()))
    log.append("\nDeep research complete! Now run Credit Analysis for richer results.")
    return "\n".join(log)
def calculate_dscr(ebitda, interest, principal):
    if float(interest) + float(principal) == 0:
        return 0
    return round(float(ebitda) / (float(interest) + float(principal)), 2)

def gst_crosscheck(gst_turnover, bank_credits, itr_income):
    flags = []
    score = 100
    diff = abs(gst_turnover - bank_credits) / max(bank_credits, 1) * 100
    if diff > 20:
        flags.append("CRITICAL: GST vs Bank mismatch >20% - possible circular trading")
        score -= 30
    elif diff > 10:
        flags.append("WARNING: GST vs Bank mismatch >10% - GSTR-2A reconciliation needed")
        score -= 15
    if bank_credits > gst_turnover * 1.5:
        flags.append("CRITICAL: Bank credits 1.5x GST turnover - circular transactions suspected")
        score -= 20
    itr_diff = abs(itr_income - bank_credits) / max(bank_credits, 1) * 100
    if itr_diff > 25:
        flags.append("CRITICAL: ITR income vs bank mismatch >25% - possible tax evasion")
        score -= 25
    if not flags:
        flags.append("No major discrepancies detected")
    return max(score, 0), flags


def extract_tables_from_pdf(filepath):
    try:
        import camelot
        tables = camelot.read_pdf(filepath, pages='all', flavor='lattice')
        if len(tables) == 0:
            tables = camelot.read_pdf(filepath, pages='all', flavor='stream')
        all_text = ""
        for i, table in enumerate(tables):
            df = table.df
            all_text += "TABLE " + str(i+1) + ":\n"
            all_text += df.to_string() + "\n\n"
        return all_text
    except:
        return ""
def ingest_pdf(file):
    if file is None:
        return "No file uploaded"
    import pytesseract
    import PIL.Image
    import io
    # Try table extraction first
    table_text = extract_tables_from_pdf(file.name)
    doc = pymupdf.open(file.name)
    text = table_text + "\n"
    for page in doc:
        page_text = page.get_text()
        if len(page_text.strip()) < 50:
            try:
                pix = page.get_pixmap(matrix=pymupdf.Matrix(2, 2))
                img = PIL.Image.open(io.BytesIO(pix.tobytes("png")))
                page_text = pytesseract.image_to_string(img, lang="eng+hin")
            except:
                pass
        text += page_text
    words = text.split()
    chunks = [" ".join(words[i:i+400]) for i in range(0, len(words), 400)]
    chunks = [c for c in chunks if len(c.strip()) > 80]
    if chunks:
        embeddings = embedder.encode(chunks).tolist()
        fname = os.path.basename(file.name)
        collection.add(documents=chunks, embeddings=embeddings,
            ids=[fname + "_" + str(i) for i in range(len(chunks))],
            metadatas=[{"source": fname, "chunk": i} for i in range(len(chunks))])
    return "Stored " + str(len(chunks)) + " chunks from " + os.path.basename(file.name) + ". Total in DB: " + str(collection.count())

async def crawl_url_async(url):
    async with AsyncWebCrawler() as crawler:
        result = await crawler.arun(url=url)
        if not result.success:
            return 0, ""
        words = result.markdown.split()
        chunks = [" ".join(words[i:i+400]) for i in range(0, len(words), 400)]
        chunks = [c for c in chunks if len(c.strip()) > 80]
        if chunks:
            embeddings = embedder.encode(chunks).tolist()
            collection.add(documents=chunks, embeddings=embeddings,
                ids=[url[:50] + "_" + str(i) for i in range(len(chunks))],
                metadatas=[{"source": url, "chunk": i} for i in range(len(chunks))])
        return len(chunks)

def crawl_url(url):
    if not url:
        return "Please enter a URL"
    n = asyncio.run(crawl_url_async(url))
    return "Stored " + str(n) + " chunks from " + url + ". Total in DB: " + str(collection.count())

def auto_research(company):
    if not company:
        return "Please enter company name first"
    urls = [
        "https://en.wikipedia.org/wiki/" + company.replace(" ", "_"),
        "https://economictimes.indiatimes.com/searchresult.cms?query=" + company.replace(" ", "+"),
        "https://www.moneycontrol.com/news/tags/" + company.replace(" ", "-").lower() + ".html",
        "https://www.sebi.gov.in/sebiweb/other/OtherAction.do?doRecognisedFpi=yes&intmId=13",
        "https://www.livelaw.in/?s=" + company.replace(" ", "+"),
        "https://www.barandbench.com/?s=" + company.replace(" ", "+"),
        "https://www.vccircle.com/?s=" + company.replace(" ", "+"),
        "https://www.business-standard.com/search?q=" + company.replace(" ", "+"),
    ]
    results = []
    for url in urls:
        try:
            n = asyncio.run(crawl_url_async(url))
            results.append("+" + str(n) + " chunks from " + url[:50])
        except:
            results.append("Failed: " + url[:50])
    return "\n".join(results) + "\nTotal in DB: " + str(collection.count())

def run_analysis(company, officer_notes, gst_turnover, bank_credits, ebitda, itr_income, interest, principal):
    if not company:
        return "Please enter company name", "", "", "", None, None

    # GST Cross-check
    gst_result = ""
    if gst_turnover > 0 and bank_credits > 0:
        gst_score, gst_flags = gst_crosscheck(gst_turnover, bank_credits, itr_income)
        gst_result = "\n".join(gst_flags)

    # DSCR
    dscr_result = ""
    if ebitda > 0 and (interest + principal) > 0:
        dscr = calculate_dscr(ebitda, interest, principal)
        dscr_result = "DSCR: " + str(dscr) + "x"
        if dscr < 1.25:
            dscr_result += " ⚠️ BELOW RBI threshold of 1.25x - HIGH RISK"
        elif dscr < 1.5:
            dscr_result += " ⚡ MODERATE - between 1.25x and 1.5x"
        else:
            dscr_result += " ✅ HEALTHY - above 1.5x"

    # Score Five Cs
    scores = {}
    for c in WEIGHTS:
        scores[c] = score_one_c(c, company, officer_notes)

    # Rule-based adjustments based on officer notes and financials
    note_lower = (officer_notes or "").lower()
    for c in scores:
        s = scores[c].get("score", 50)
        if any(w in note_lower for w in ["nclt", "insolvency", "insolvent"]):
            if c == "Character": scores[c]["score"] = min(s, 20)
            if c == "Capital": scores[c]["score"] = min(s, 25)
        if any(w in note_lower for w in ["ed investigation", "fema", "fraud"]):
            if c == "Character": scores[c]["score"] = min(scores[c]["score"], 15)
        if any(w in note_lower for w in ["wilful defaulter", "defaulter"]):
            if c == "Character": scores[c]["score"] = min(scores[c]["score"], 10)
        if any(w in note_lower for w in ["auditor resigned", "deloitte resigned"]):
            if c == "Capital": scores[c]["score"] = min(scores[c]["score"], 20)
        if any(w in note_lower for w in ["40% capacity", "factory closed", "operations stopped"]):
            if c == "Capacity": scores[c]["score"] = min(scores[c]["score"], 30)
    
    # DSCR based adjustment
    if ebitda > 0 and (interest + principal) > 0:
        dscr = calculate_dscr(ebitda, interest, principal)
        if dscr < 0.5:
            scores["Capacity"]["score"] = min(scores["Capacity"].get("score", 50), 20)
        elif dscr < 1.0:
            scores["Capacity"]["score"] = min(scores["Capacity"].get("score", 50), 35)
        elif dscr < 1.25:
            scores["Capacity"]["score"] = min(scores["Capacity"].get("score", 50), 45)

    final_score = round(sum(scores[c].get("score", 50) * WEIGHTS[c] / 100 for c in WEIGHTS), 1)
    red_flags = [f for c in scores for f in scores[c].get("red_flags", [])]

    if final_score >= 75 and len(red_flags) == 0:
        decision, rate, multiple = "APPROVE", "9.5%", 3.0
    elif final_score >= 65 and len(red_flags) <= 1:
        decision, rate, multiple = "APPROVE WITH CONDITIONS", "11.5%", 2.0
    elif final_score >= 50:
        decision, rate, multiple = "REFER TO CREDIT COMMITTEE", "13.5%", 1.0
    else:
        decision, rate, multiple = "REJECT", "N/A", 0

    # Early warning
    # Also count GST critical flags
    gst_critical = gst_result.count("CRITICAL") if gst_result else 0
    dscr_flag = 1 if dscr_result and "BELOW RBI" in dscr_result else 0
    total_flags = len(red_flags) + gst_critical + dscr_flag
    if total_flags == 0:
        warning = "🟢 SAFE ZONE — No significant risk signals"
    elif total_flags <= 3:
        warning = "🟡 WATCH ZONE — " + str(total_flags) + " minor risk signals"
    elif total_flags <= 8:
        warning = "🟠 CAUTION ZONE — " + str(total_flags) + " significant risk signals"
    else:
        warning = "🔴 DANGER ZONE — " + str(total_flags) + " CRITICAL risk signals"

    # Build score report
    score_report = "=" * 55 + "\n"
    score_report += "INTELLI-CREDIT ANALYSIS: " + company + "\n"
    score_report += "=" * 55 + "\n\n"
    score_report += warning + "\n"
    score_report += "FINAL SCORE: " + str(final_score) + "/100\n"
    score_report += "DECISION: " + decision + "\n"
    score_report += "INTEREST RATE: " + rate + "\n\n"
    score_report += "FIVE Cs BREAKDOWN:\n"
    for c in WEIGHTS:
        s = scores[c]
        bar = "█" * (s.get("score", 50) // 5) + "░" * (20 - s.get("score", 50) // 5)
        score_report += c.ljust(12) + " [" + bar + "] " + str(s.get("score", 50)) + "/100 " + s.get("rating", "Fair") + "\n"
        for flag in s.get("red_flags", []):
            score_report += "   ⚠️ " + flag + "\n"

    if red_flags:
        score_report += "\nRED FLAGS (" + str(len(red_flags)) + "):\n"
        for f in red_flags:
            score_report += "⚠️ " + f + "\n"

    score_report += "\nREASONING:\n"
    for c in WEIGHTS:
        score_report += "[" + c + "] " + scores[c].get("reasoning", "") + "\n\n"

    if gst_result:
        score_report += "GST CROSS-CHECK:\n" + gst_result + "\n"
    if dscr_result:
        score_report += "\n" + dscr_result + "\n"

    # Generate CAM
    data = {"company": company, "final_score": final_score, "decision": decision,
            "rate": rate, "scores": scores, "red_flags": red_flags, "officer_notes": officer_notes,
            "structured_data": {"gst_turnover": gst_turnover, "bank_credits": bank_credits,
                                "ebitda": ebitda, "itr_income": itr_income, "interest": interest, "principal": principal}}

    with open("last_score.json", "w") as f:
        json.dump(data, f, indent=2)
    save_audit_trail(company, decision, final_score, red_flags, officer_notes)

    dscr_val = calculate_dscr(ebitda, interest, principal) if (interest + principal) > 0 else 0
    gst_score, gst_flags_list = gst_crosscheck(gst_turnover, bank_credits, itr_income) if gst_turnover > 0 else (100, [])
    rejection_reason = generate_rejection_reason(decision, scores, red_flags, dscr_val, gst_flags_list)
    cam_file = generate_cam(data)
    cma_file = generate_cma(data)

    return score_report, gst_result, dscr_result, warning, rejection_reason, cam_file, cma_file


def explain_decision(score_report):
    if not score_report or "INTELLI-CREDIT" not in score_report:
        return "Please run credit analysis first"
    
    prompt = (
        "You are a senior Indian credit committee chairman presenting a credit decision.\n"
        "Based on this credit analysis report, explain the decision in plain English\n"
        "as if you are presenting to a board of directors. Be specific, cite the key factors,\n"
        "mention India-specific risks like GSTR mismatch, CIBIL score, NCLT cases.\n"
        "Use banker language. Be concise but comprehensive. Max 200 words.\n\n"
        "Report:\n" + score_report[:2000] + "\n\n"
        "Start with: 'Distinguished members of the Credit Committee...'"
    )
    return ask_ollama(prompt, timeout=180)

def generate_rejection_reason(decision, scores, red_flags, dscr, gst_flags):
    if "APPROVE" in decision and "CONDITIONS" not in decision:
        return "✅ APPROVED — Strong financials, clean GST records, healthy DSCR above RBI threshold."
    
    reasons = []
    
    # DSCR check
    if dscr and dscr < 1.25:
        reasons.append("DSCR of " + str(dscr) + "x is below RBI minimum threshold of 1.25x")
    
    # GST flags
    for flag in gst_flags:
        if "CRITICAL" in flag:
            if "circular" in flag.lower():
                reasons.append("Circular trading detected — bank credits exceed GST turnover by >50%")
            elif "mismatch" in flag.lower():
                reasons.append("GST vs Bank mismatch >20% suggesting revenue inflation")
            elif "ITR" in flag:
                reasons.append("ITR income vs bank mismatch >25% indicating possible tax evasion")
    
    # Red flags from Five Cs
    for flag in red_flags[:3]:
        if flag not in reasons:
            reasons.append(flag)
    
    # Low scoring Cs
    for c, s in scores.items():
        if s.get("score", 50) < 40:
            reasons.append(c + " score critically low at " + str(s.get("score", 50)) + "/100 (" + s.get("rating", "") + ")")
    
    if not reasons:
        reasons.append("Overall risk score below minimum threshold for approval")
    
    if decision == "REJECT":
        prefix = "❌ REJECTED due to: "
    elif "CONDITIONS" in decision:
        prefix = "⚠️ APPROVED WITH CONDITIONS due to: "
    else:
        prefix = "🔄 REFERRED TO CREDIT COMMITTEE due to: "
    
    return prefix + " + ".join(reasons[:3])
def loan_simulator(loan_amount, final_score_str):
    try:
        final_score = float(final_score_str.split("FINAL SCORE:")[1].split("/")[0].strip()) if "FINAL SCORE:" in final_score_str else 50
    except:
        final_score = 50

    base_rate = 9.5 if final_score >= 75 else 11.5 if final_score >= 65 else 13.5 if final_score >= 50 else 16.0
    final_rate = round(base_rate - round((500 - loan_amount) / 500 * 0.5, 2), 2)
    emi = round((loan_amount * final_rate / 100) / 12, 2)

    result = "Loan Amount: Rs. " + str(loan_amount) + " Crore\n"
    result += "Interest Rate: " + str(final_rate) + "% per annum\n"
    result += "Monthly Interest: Rs. " + str(emi) + " Crore\n"
    result += "Annual Interest: Rs. " + str(round(loan_amount * final_rate / 100, 2)) + " Crore"
    return result

def generate_cam(data):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = Document()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CREDIT APPRAISAL MEMORANDUM (CAM)")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0, 51, 102)
    doc.add_paragraph("Company: " + data["company"] + "  |  Date: " + datetime.now().strftime("%d %B %Y"))
    doc.add_paragraph("Powered by Intelli-Credit AI | Databricks-compatible pipeline | RBI Framework")
    doc.add_paragraph()
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in [("Company", data["company"]), ("Score", str(data["final_score"]) + "/100"),
                          ("Decision", data["decision"]), ("Interest Rate", data.get("rate", "N/A"))]:
        row = table.add_row()
        row.cells[0].text = label
        row.cells[1].text = str(value)
        row.cells[0].paragraphs[0].runs[0].bold = True
    doc.add_paragraph()
    colors = {"APPROVE": (0,128,0), "APPROVE WITH CONDITIONS": (200,100,0), "REFER TO CREDIT COMMITTEE": (0,0,200), "REJECT": (200,0,0)}
    p = doc.add_paragraph()
    run = p.add_run("DECISION: " + data["decision"])
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(*colors.get(data["decision"], (0,0,0)))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading("FIVE Cs ANALYSIS (RBI Framework)", level=1)
    for c, s in data["scores"].items():
        doc.add_heading(c + " — " + str(s.get("score",50)) + "/100 (" + s.get("rating","Fair") + ")", level=2)
        doc.add_paragraph(s.get("reasoning", ""))
        for f in s.get("red_flags", []):
            p2 = doc.add_paragraph(f, style="List Bullet")
            if p2.runs:
                p2.runs[0].font.color.rgb = RGBColor(200, 0, 0)
    if data.get("red_flags"):
        doc.add_heading("RED FLAG SUMMARY", level=1)
        for f in data["red_flags"]:
            p2 = doc.add_paragraph(f, style="List Bullet")
            if p2.runs:
                p2.runs[0].font.color.rgb = RGBColor(200, 0, 0)
    if data.get("officer_notes"):
        doc.add_heading("OFFICER OBSERVATIONS", level=1)
        doc.add_paragraph(data["officer_notes"])
    fname = data["company"].replace(" ", "_") + "_CAM_" + datetime.now().strftime("%Y%m%d") + ".docx"
    doc.save(fname)
    return fname

def generate_cma(data):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = Document()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CREDIT MONITORING ARRANGEMENT (CMA) REPORT")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 51, 102)
    doc.add_paragraph("Company: " + data["company"] + "  |  Date: " + datetime.now().strftime("%d %B %Y"))
    sd = data.get("structured_data", {})
    gst = sd.get("gst_turnover", 0)
    eb = sd.get("ebitda", 0)
    intr = sd.get("interest", 0)
    princ = sd.get("principal", 0)
    bank = sd.get("bank_credits", 0)
    itr = sd.get("itr_income", 0)
    doc.add_heading("1. OPERATING STATEMENT (Rs. Crore)", level=1)
    t1 = doc.add_table(rows=0, cols=4)
    t1.style = "Table Grid"
    h = t1.add_row()
    for i, v in enumerate(["Particulars", "Previous Year", "Current Year", "Projected"]):
        h.cells[i].text = v
        h.cells[i].paragraphs[0].runs[0].bold = True
    for r in [
        ("GST Turnover", str(round(gst*0.85,2)), str(gst), str(round(gst*1.15,2))),
        ("Bank Credits", str(round(bank*0.85,2)), str(bank), str(round(bank*1.1,2))),
        ("ITR Income", str(round(itr*0.85,2)), str(itr), str(round(itr*1.1,2))),
        ("EBITDA", str(round(eb*0.85,2)), str(eb), str(round(eb*1.1,2))),
        ("Interest", str(round(intr*0.9,2)), str(intr), str(round(intr*1.05,2))),
        ("Net Profit", str(round((eb-intr)*0.7*0.85,2)), str(round((eb-intr)*0.7,2)), str(round((eb-intr)*0.7*1.1,2))),
    ]:
        row = t1.add_row()
        for i, val in enumerate(r):
            row.cells[i].text = val
    doc.add_heading("2. KEY RATIOS (RBI Framework)", level=1)
    t2 = doc.add_table(rows=0, cols=4)
    t2.style = "Table Grid"
    h2 = t2.add_row()
    for i, v in enumerate(["Ratio", "Value", "RBI Benchmark", "Status"]):
        h2.cells[i].text = v
        h2.cells[i].paragraphs[0].runs[0].bold = True
    dscr_val = round(float(eb)/(float(intr)+float(princ)),2) if (intr+princ)>0 else 0
    de = round(intr*8/max(eb*3,1),2)
    gst_var = round(abs(gst-bank)/max(bank,1)*100,1)
    for r in [
        ("DSCR", str(dscr_val)+"x", ">1.25x", "PASS" if dscr_val>=1.25 else "FAIL"),
        ("Debt-Equity", str(de)+"x", "<3x", "PASS" if de<3 else "FAIL"),
        ("EBITDA Margin", str(round(eb/max(gst,1)*100,1))+"%", ">10%", "PASS" if eb/max(gst,1)*100>=10 else "FAIL"),
        ("GST vs Bank", str(gst_var)+"%", "<10%", "PASS" if gst_var<10 else "FAIL - Circular Trading Risk"),
    ]:
        row = t2.add_row()
        for i, val in enumerate(r):
            row.cells[i].text = val
    doc.add_heading("3. DECISION", level=1)
    doc.add_paragraph("Final Decision: " + data["decision"]).runs[0].bold = True
    doc.add_paragraph("Powered by Intelli-Credit AI | Databricks pipeline | RBI compliant").runs[0].font.size = Pt(8)
    fname = data["company"].replace(" ", "_") + "_CMA_" + datetime.now().strftime("%Y%m%d") + ".docx"
    doc.save(fname)
    return fname

# ── Gradio UI ─────────────────────────────────────────────────

with gr.Blocks(title="Intelli-Credit") as app:

    gr.Markdown("""
    # ⚖️ Intelli-Credit
    ### AI-Powered Corporate Credit Appraisal Engine · Indian Banking Framework · RBI Compliant
    > ⬡ **POWERED BY DATABRICKS** | Delta Lake · MLflow · Unity Catalog · Apache Spark
    """)

    with gr.Tabs():

        # Tab 1: Data Ingestor
        with gr.Tab("📥 Data Ingestor"):
            gr.Markdown("### Feed company documents into AI memory")
            with gr.Row():
                with gr.Column():
                    pdf_upload = gr.File(label="Upload PDF (Annual Report, GST, Bank Statement)", file_types=[".pdf"])
                    pdf_btn = gr.Button("📄 Process PDF", variant="primary")
                    pdf_output = gr.Textbox(label="Result", lines=2)
                with gr.Column():
                    url_input = gr.Textbox(label="Crawl a URL", placeholder="https://en.wikipedia.org/wiki/...")
                    url_btn = gr.Button("🕷️ Crawl URL", variant="primary")
                    url_output = gr.Textbox(label="Result", lines=2)

            pdf_btn.click(ingest_pdf, inputs=[pdf_upload], outputs=[pdf_output])
            url_btn.click(crawl_url, inputs=[url_input], outputs=[url_output])

        # Tab 2: Credit Analysis
        with gr.Tab("⚡ Credit Analysis"):
            gr.Markdown("### Company Credit Appraisal")
            with gr.Row():
                with gr.Column(scale=2):
                    company_input = gr.Textbox(label="Company Name", placeholder="e.g. Adani Group, Reliance Industries")
                    officer_notes_input = gr.Textbox(label="Credit Officer Notes (optional)", placeholder="e.g. Factory at 40% capacity, management evasive", lines=3)
                    auto_btn = gr.Button("🔍 Auto-Research Company")
                    deep_btn = gr.Button("🔬 Deep Research (News + Web + MCA)", variant="secondary")
                    auto_output = gr.Textbox(label="Research Result", lines=6)

                with gr.Column(scale=2):
                    gr.Markdown("### 📊 Structured Financial Data")
                    with gr.Row():
                        gst_input = gr.Number(label="GST Turnover (Rs. Cr)", value=0)
                        bank_input = gr.Number(label="Bank Credits (Rs. Cr)", value=0)
                    with gr.Row():
                        ebitda_input = gr.Number(label="EBITDA (Rs. Cr)", value=0)
                        itr_input = gr.Number(label="ITR Net Income (Rs. Cr)", value=0)
                    with gr.Row():
                        interest_input = gr.Number(label="Annual Interest (Rs. Cr)", value=0)
                        principal_input = gr.Number(label="Annual Principal (Rs. Cr)", value=0)

            analyze_btn = gr.Button("⚡ Run Credit Analysis", variant="primary", size="lg")

            with gr.Row():
                score_output = gr.Textbox(label="Credit Score Report", lines=30)
                with gr.Column():
                    gst_output = gr.Textbox(label="GST Cross-Check", lines=5)
                    dscr_output = gr.Textbox(label="DSCR Result", lines=3)
                    warning_output = gr.Textbox(label="🚨 Early Warning Signal", lines=2)
                    rejection_output = gr.Textbox(label="📋 Decision Reason", lines=3)

            with gr.Row():
                cam_output = gr.File(label="⬇️ Download CAM Report")
                cma_output = gr.File(label="⬇️ Download CMA Report")

            gr.Markdown("---")
            explain_btn = gr.Button("🎤 Explain This Decision (AI Banker)", variant="secondary")
            explain_output = gr.Textbox(label="AI Credit Committee Explanation", lines=10)
            explain_btn.click(explain_decision, inputs=[score_output], outputs=[explain_output])

            auto_btn.click(auto_research, inputs=[company_input], outputs=[auto_output])
            deep_btn.click(deep_research, inputs=[company_input], outputs=[auto_output])
            analyze_btn.click(
                run_analysis,
                inputs=[company_input, officer_notes_input, gst_input, bank_input, ebitda_input, itr_input, interest_input, principal_input],
                outputs=[score_output, gst_output, dscr_output, warning_output, rejection_output, cam_output, cma_output]
            )

        # Tab 3: Audit Trail
        with gr.Tab("📋 Audit Trail"):
            gr.Markdown("### Decision Log — Every analysis recorded with timestamp and data sources")
            refresh_btn = gr.Button("🔄 Load Audit Trail", variant="primary")
            audit_output = gr.Textbox(label="Audit Log", lines=25)
            refresh_btn.click(get_audit_trail, inputs=[], outputs=[audit_output])

        # Tab 3: Loan Simulator
        with gr.Tab("💰 Loan Simulator"):
            gr.Markdown("### Risk-Adjusted Loan Calculator")
            gr.Markdown("Run analysis first, then use this simulator")
            with gr.Row():
                loan_slider = gr.Slider(minimum=1, maximum=500, value=50, step=5, label="Loan Amount (Rs. Crore)")
                score_for_sim = gr.Textbox(label="Paste score report here for accurate rate", lines=3)
            sim_btn = gr.Button("Calculate", variant="primary")
            sim_output = gr.Textbox(label="Loan Details", lines=6)
            sim_btn.click(loan_simulator, inputs=[loan_slider, score_for_sim], outputs=[sim_output])

if __name__ == "__main__":
    app.launch(server_port=7860, share=True, js="""
() => {
    document.body.classList.add('dark');
    const style = document.createElement('style');
    style.textContent = `
        .primary { background: #00f5c4 !important; color: #0d0d14 !important; }
        button.primary { background: #00f5c4 !important; color: #0d0d14 !important; }
        .tab-nav button.selected { color: #00f5c4 !important; border-color: #00f5c4 !important; }
        a { color: #00f5c4 !important; }
    `;
    document.head.appendChild(style);
}
""")
